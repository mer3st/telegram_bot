import os
import logging

logger = logging.getLogger(__name__)


def detect_ext(path: str) -> str:
    """Detect file extension from magic bytes when filename has none."""
    try:
        with open(path, "rb") as f:
            header = f.read(8)
        if header[:4] == b"PK\x03\x04":
            import zipfile
            with zipfile.ZipFile(path) as z:
                names = z.namelist()
            if any(n.startswith("word/") for n in names):
                return ".docx"
            if any(n.startswith("xl/") for n in names):
                return ".xlsx"
            return ".docx"
        if header[:4] == b"%PDF":
            return ".pdf"
    except Exception:
        pass
    return ""


def extract_text(path: str) -> str:
    """Extract plain text from .xlsx/.docx/.pdf/.zip/.txt files. Returns '' on failure."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".xlsx":
            return _xlsx(path)
        if ext == ".docx":
            return _docx(path)
        if ext == ".pdf":
            return _pdf(path)
        if ext == ".zip":
            return _zip(path)
        if ext in (".txt", ".csv"):
            return _txt(path)
    except Exception as e:
        logger.warning("extract_text failed for %s: %s", path, e)
    return ""


def _xlsx(path: str) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(path, data_only=True, read_only=True)
    lines = []
    for ws in wb.worksheets:
        lines.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lines.append("\t".join(cells))
    wb.close()
    return "\n".join(lines)


_WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _docx(path: str) -> str:
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    # Extract text from text boxes using full Clark notation (wps prefix not in python-docx registry)
    for el in d.element.body.iter(f"{{{_WPS}}}txbx"):
        for p in el.iter(f"{{{_W}}}p"):
            runs = [r.text for r in p.iter(f"{{{_W}}}t") if r.text]
            if runs:
                parts.append("".join(runs))
    return "\n".join(parts)


def _pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _zip(path: str) -> str:
    import zipfile, tempfile
    parts = []
    with zipfile.ZipFile(path, "r") as z:
        for name in z.namelist():
            ext = os.path.splitext(name)[1].lower()
            if ext not in (".xlsx", ".docx", ".pdf", ".txt", ".csv"):
                continue
            tmp_path = None  # #10: declare before try so finally can always reference it
            try:
                with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                    tmp.write(z.read(name))
                    tmp_path = tmp.name
                txt = extract_text(tmp_path)
                if txt.strip():
                    parts.append(f"=== {name} ===\n{txt}")
            except Exception as e:
                logger.warning("zip entry %s failed: %s", name, e)
            finally:
                if tmp_path and os.path.exists(tmp_path):
                    os.remove(tmp_path)
    return "\n\n".join(parts)